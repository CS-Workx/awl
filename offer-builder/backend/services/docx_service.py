from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Dict, Any
import os
from datetime import datetime
import uuid
import re


class DocxService:
    def __init__(self, template_path: str, output_dir: str):
        self.template_path = template_path
        self.output_dir = output_dir
        
        # Ensure output directory exists
        os.makedirs(output_dir, exist_ok=True)
    
    def create_offer_docx(
        self,
        offer_data: Dict[str, Any],
        crm_data: Dict[str, Any],
        training_data: Dict[str, Any],
        intake_data: Dict[str, Any]
    ) -> str:
        """
        Generate formatted DOCX from template with offer data
        """
        try:
            # Check if template exists
            if self.template_path and os.path.exists(self.template_path):
                # Load template
                doc = Document(self.template_path)
                
                # Build comprehensive field map
                field_map = self._build_field_map(offer_data, crm_data, training_data, intake_data)
                
                # Replace all placeholders using new method
                self._find_and_replace_fields(doc, field_map)
                
                # Update pricing table if available
                if len(doc.tables) >= 1:
                    self._update_pricing_table(doc.tables[0], offer_data, training_data, intake_data)
            else:
                # Create new document from scratch
                doc = self._create_new_document(offer_data, crm_data, training_data, intake_data)
            
            # Generate filename with new format
            filename = self._generate_filename(crm_data, training_data)
            filepath = os.path.join(self.output_dir, filename)
            
            # Save document
            doc.save(filepath)
            
            return filepath
            
        except Exception as e:
            print(f"Error creating DOCX: {str(e)}")
            raise Exception(f"Failed to create DOCX: {str(e)}")
    
    def _build_field_map(
        self,
        offer_data: Dict[str, Any],
        crm_data: Dict[str, Any],
        training_data: Dict[str, Any],
        intake_data: Dict[str, Any]
    ) -> Dict[str, str]:
        """Build complete mapping of all possible template fields"""
        
        # Generate reference number with new format
        ref_number = self._generate_offer_number(crm_data, training_data)
        
        # Format date in Dutch
        offer_date = self._format_dutch_date(crm_data.get('datum_offerte'))
        
        return {
            # Reference & Dates
            'reference_number': ref_number,
            'referentie': ref_number,
            'offer_date': offer_date,
            'datum_offerte': offer_date,
            'datum': offer_date,
            
            # Client Info
            'client_name': crm_data.get('potentiele_klant', ''),
            'potentiele_klant': crm_data.get('potentiele_klant', ''),
            'client_company_address': crm_data.get('client_company_address', ''),
            'client_vat_number': crm_data.get('client_vat_number', ''),
            'primaire_contact': crm_data.get('primaire_contact', ''),
            'extra_contact': crm_data.get('extra_contact', ''),
            
            # Advisor
            'advisor_name': crm_data.get('advisor_name', crm_data.get('type', '')),
            'advisor_phone': crm_data.get('advisor_phone', ''),
            'advisor_email': crm_data.get('advisor_email', ''),
            
            # Project Manager
            'project_manager_name': crm_data.get('project_manager_name', ''),
            'project_manager_phone': crm_data.get('project_manager_phone', ''),
            'project_manager_email': crm_data.get('project_manager_email', ''),
            
            # Training Info
            'training_title': training_data.get('title', ''),
            'training_duration': training_data.get('duration', ''),
            'training_objectives': offer_data.get('training_objectives', training_data.get('learning_objectives', '')),
            'program_outline': offer_data.get('program_overview', training_data.get('program_outline', '')),
            'number_of_participants': str(intake_data.get('number_of_participants', '')),
            
            # Offer Content
            'introduction': offer_data.get('introduction', ''),
            'practical_arrangements': offer_data.get('practical_arrangements', ''),
            'investment': offer_data.get('investment', ''),
            'next_steps': offer_data.get('next_steps', ''),
        }
    
    def _format_dutch_date(self, date_str: str = None) -> str:
        """Format date in Dutch: 'DD month YYYY'"""
        if not date_str:
            date_str = datetime.now().strftime('%d/%m/%Y')
        
        try:
            # Try to parse DD/MM/YYYY format
            date_obj = datetime.strptime(date_str, '%d/%m/%Y')
        except:
            try:
                # Try YYYY-MM-DD format
                date_obj = datetime.strptime(date_str, '%Y-%m-%d')
            except:
                # Return as-is if can't parse
                return date_str
        
        months_nl = ['januari', 'februari', 'maart', 'april', 'mei', 'juni',
                     'juli', 'augustus', 'september', 'oktober', 'november', 'december']
        return f"{date_obj.day} {months_nl[date_obj.month - 1]} {date_obj.year}"
    
    def _find_and_replace_fields(self, doc: Document, field_map: Dict[str, str]):
        """
        Find placeholder patterns in template and replace with actual values
        Patterns: {{field_name}} or [FIELD_NAME] or specific Dutch labels
        """
        for para in doc.paragraphs:
            for field_key, field_value in field_map.items():
                if not field_value:
                    field_value = ''
                
                # Support multiple placeholder formats
                patterns = [
                    f"{{{{{field_key}}}}}",  # {{client_name}}
                    f"[{field_key.upper()}]",  # [CLIENT_NAME]
                ]
                
                for pattern in patterns:
                    if pattern in para.text:
                        self._replace_text_in_runs(para, pattern, field_value)
        
        # Also search in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for field_key, field_value in field_map.items():
                            if not field_value:
                                field_value = ''
                            
                            patterns = [
                                f"{{{{{field_key}}}}}",
                                f"[{field_key.upper()}]",
                            ]
                            
                            for pattern in patterns:
                                if pattern in para.text:
                                    self._replace_text_in_runs(para, pattern, field_value)
    
    def _contains_markdown(self, text: str) -> bool:
        """
        Check if text contains Markdown formatting
        Detects: **bold**, *italic*, - bullets, 1. numbered lists, # headings
        """
        if not text:
            return False
        
        markdown_patterns = [
            r'\*\*[^*]+\*\*',          # **bold**
            r'(?<!\*)\*(?!\*)([^*]+)\*(?!\*)',  # *italic* (not part of **)
            r'^[\-\*\+]\s+',           # - bullet or * bullet at line start
            r'^\d+\.\s+',              # 1. numbered list at line start
            r'^#{1,6}\s+',             # # heading at line start
        ]
        
        for pattern in markdown_patterns:
            if re.search(pattern, text, re.MULTILINE):
                return True
        return False
    
    def _add_markdown_text(self, paragraph, markdown_text: str):
        """
        Add text with inline Markdown formatting to paragraph
        Handles: **bold**, *italic*, ***bold+italic***
        """
        if not markdown_text:
            return
        
        # Pattern to match ***bold+italic***, **bold**, or *italic*
        # Must be greedy for *** first, then **, then *
        pattern = r'(\*\*\*(?:[^*]|\*(?!\*\*))+\*\*\*|\*\*(?:[^*]|\*(?!\*))+\*\*|\*(?:[^*])+\*)'
        
        last_pos = 0
        for match in re.finditer(pattern, markdown_text):
            # Add plain text before match
            if match.start() > last_pos:
                plain_text = markdown_text[last_pos:match.start()]
                if plain_text:
                    paragraph.add_run(plain_text)
            
            # Add formatted text
            formatted_text = match.group(0)
            
            if formatted_text.startswith('***') and formatted_text.endswith('***'):
                # Bold + Italic
                text = formatted_text[3:-3]
                run = paragraph.add_run(text)
                run.bold = True
                run.italic = True
            elif formatted_text.startswith('**') and formatted_text.endswith('**'):
                # Bold
                text = formatted_text[2:-2]
                run = paragraph.add_run(text)
                run.bold = True
            elif formatted_text.startswith('*') and formatted_text.endswith('*'):
                # Italic
                text = formatted_text[1:-1]
                run = paragraph.add_run(text)
                run.italic = True
            
            last_pos = match.end()
        
        # Add remaining plain text
        if last_pos < len(markdown_text):
            remaining = markdown_text[last_pos:]
            if remaining:
                paragraph.add_run(remaining)
    
    def _add_markdown_content(self, doc, markdown_text: str):
        """
        Parse Markdown and add formatted content to document
        Handles: headings, bullet lists, numbered lists, paragraphs with inline formatting
        """
        if not markdown_text:
            return
        
        # Split into lines
        lines = markdown_text.split('\n')
        
        i = 0
        while i < len(lines):
            line = lines[i].rstrip()
            
            # Skip empty lines
            if not line.strip():
                i += 1
                continue
            
            # Check for heading (# H1, ## H2, etc.)
            heading_match = re.match(r'^(#{1,6})\s+(.+)$', line)
            if heading_match:
                level = len(heading_match.group(1))
                text = heading_match.group(2)
                para = doc.add_heading(level=level)
                # Clear default text and add with formatting
                para.text = ''
                self._add_markdown_text(para, text)
                i += 1
                continue
            
            # Check for bullet list (-, *, +)
            bullet_match = re.match(r'^[\-\*\+]\s+(.+)$', line)
            if bullet_match:
                text = bullet_match.group(1)
                para = doc.add_paragraph(style='List Bullet')
                self._add_markdown_text(para, text)
                i += 1
                continue
            
            # Check for numbered list (1., 2., etc.)
            number_match = re.match(r'^\d+\.\s+(.+)$', line)
            if number_match:
                text = number_match.group(1)
                para = doc.add_paragraph(style='List Number')
                self._add_markdown_text(para, text)
                i += 1
                continue
            
            # Regular paragraph
            para = doc.add_paragraph()
            self._add_markdown_text(para, line)
            i += 1
    
    def _replace_text_in_runs(self, paragraph, old_text: str, new_text: str):
        """
        Replace text in paragraph runs while preserving formatting
        If new_text contains Markdown, format it appropriately
        """
        if old_text in paragraph.text:
            # Check if new_text contains Markdown
            if self._contains_markdown(new_text):
                # Clear the paragraph and add formatted content
                paragraph.clear()
                self._add_markdown_text(paragraph, new_text)
            else:
                # Simple text replacement
                for run in paragraph.runs:
                    if old_text in run.text:
                        run.text = run.text.replace(old_text, new_text)
    
    def _replace_placeholders(
        self,
        doc: Document,
        offer_data: Dict[str, Any],
        crm_data: Dict[str, Any],
        training_data: Dict[str, Any],
        intake_data: Dict[str, Any]
    ):
        """
        Replace content in Syntra Bizz template document
        The template has specific sections that need to be filled in
        """
        # Find and fill specific sections in the template
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            
            # Fill "Referentie" section (paragraph index ~5)
            if 'Referentie' in text and len(text) < 50:
                if i + 1 < len(doc.paragraphs):
                    doc.paragraphs[i + 1].text = self._generate_offer_number(crm_data)
            
            # Fill "Voorstel voor Opleiding" section (paragraph index ~10)
            elif 'Voorstel voor' in text and 'Opleiding' in text:
                if i + 1 < len(doc.paragraphs):
                    client_name = crm_data.get('potentiele_klant', '')
                    training_title = training_data.get('title', '')
                    # Combine on same line with separator
                    doc.paragraphs[i + 1].text = f"{client_name} - {training_title}" if training_title else client_name
            
            # Fill "Doelstelling van de opleiding" section
            elif 'Doelstelling van de opleiding' in text:
                # Insert objectives after this header
                if i + 2 < len(doc.paragraphs):
                    # Find "Aantal personen" to insert before it
                    objectives_text = offer_data.get('training_objectives', training_data.get('learning_objectives', ''))
                    if objectives_text:
                        doc.paragraphs[i + 1].text = objectives_text
            
            # Fill "Aantal personen" section
            elif 'Aantal personen:' in text and len(text) < 50:
                if i + 1 < len(doc.paragraphs):
                    num_participants = intake_data.get('number_of_participants', '')
                    doc.paragraphs[i + 1].text = str(num_participants) if num_participants else ''
            
            # Fill "Programma" section (paragraph index ~52)
            elif text == 'Programma' or text.startswith('Programma'):
                if i + 1 < len(doc.paragraphs):
                    program = offer_data.get('program_overview', training_data.get('program_outline', ''))
                    if program:
                        doc.paragraphs[i + 1].text = program
            
            # Fill "Duur van de opleiding" section
            elif 'Duur van de opleiding' in text:
                if i + 1 < len(doc.paragraphs):
                    duration = training_data.get('duration', '')
                    practical = offer_data.get('practical_arrangements', '')
                    # Combine duration and practical info on same line
                    if duration and practical:
                        combined = f"{duration}. {practical}"
                    else:
                        combined = duration or practical
                    if combined:
                        doc.paragraphs[i + 1].text = combined
            
            # Fill "Datum" at top
            elif text == 'Datum' and i < 10:
                if i + 1 < len(doc.paragraphs):
                    offer_date = crm_data.get('datum_offerte', datetime.now().strftime('%d/%m/%Y'))
                    doc.paragraphs[i + 1].text = offer_date
        
        # Fill signature page fields at the end of document
        self._fill_signature_page(doc, crm_data)
        
        # Update pricing table if available (Table 1)
        if len(doc.tables) >= 1:
            self._update_pricing_table(doc.tables[0], offer_data, training_data, intake_data)
    
    def _update_pricing_table(self, table, offer_data: Dict[str, Any], training_data: Dict[str, Any], intake_data: Dict[str, Any]):
        """
        Update the pricing table in the template
        Table structure:
        Row 0: Headers (Opleidingsonderdeel, Aantal, Eenheidsprijs, Totaalprijs)
        Row 1: Intakegesprek
        Row 2: Voorbereidend werk
        Row 3: Maatopleiding
        Row 4: Intervisie/opvolging
        """
        try:
            # Extract pricing from offer_data or training_data
            investment_text = offer_data.get('investment') or training_data.get('price') or ''
            
            # Try to extract price from text (e.g., "€ 1500" or "1500 euro")
            import re
            price_match = re.search(r'€\s*(\d+(?:[.,]\d+)?)|(\d+(?:[.,]\d+)?)\s*€', str(investment_text))
            
            if price_match:
                price_str = price_match.group(1) or price_match.group(2)
                price = price_str.replace(',', '.')
            else:
                # No price found, skip table update
                return
            
            # Update row 3 (Maatopleiding) - this is typically the main training cost
            if len(table.rows) > 3:
                row = table.rows[3]
                if len(row.cells) >= 4:
                    # Leave Opleidingsonderdeel as is
                    # Set aantal (quantity) - typically 1 or number of days
                    duration = training_data.get('duration') or ''
                    days_match = re.search(r'(\d+)\s*dag', str(duration), re.IGNORECASE)
                    quantity = days_match.group(1) if days_match else '1'
                    
                    row.cells[1].text = quantity  # Aantal
                    row.cells[2].text = f"{price} €"  # Eenheidsprijs
                    
                    # Calculate total
                    try:
                        total = float(price) * float(quantity)
                        row.cells[3].text = f"{total:.0f} €"  # Totaalprijs
                    except:
                        row.cells[3].text = f"{price} €"
        
        except Exception as e:
            print(f"Warning: Could not update pricing table: {str(e)}")
    
    def _fill_signature_page(self, doc: Document, crm_data: Dict[str, Any]):
        """
        Fill in client information on the signature page
        Uses semantic search instead of hardcoded paragraph indices
        """
        try:
            client_name = crm_data.get('potentiele_klant', '')
            contact_person = crm_data.get('primaire_contact', '')
            client_company_address = crm_data.get('client_company_address', '')
            client_vat_number = crm_data.get('client_vat_number', '')
            advisor_name = crm_data.get('advisor_name', crm_data.get('type', ''))
            formatted_date = self._format_dutch_date(crm_data.get('datum_offerte'))
            
            # Search through all paragraphs for signature page patterns
            for para in doc.paragraphs:
                text = para.text
                
                # "Opgemaakt te Berchem op [DATE]"
                if 'Opgemaakt te Berchem op' in text:
                    self._replace_text_in_runs(para, 'Opgemaakt te Berchem op', 
                                               f'Opgemaakt te Berchem op {formatted_date}')
                
                # "Voor Syntra Bizz BV         Voor [CLIENT_NAME]"
                if 'Voor Syntra Bizz' in text and text.count('Voor') >= 2:
                    # Find the second "Voor" and append client name
                    if client_name and 'Voor Syntra Bizz BV' in text:
                        # Keep existing tabs/formatting
                        new_text = text.replace('Voor Syntra Bizz BV', 
                                               f'Voor Syntra Bizz BV', 1)
                        # Add client name after second Voor
                        parts = new_text.split('Voor', 2)
                        if len(parts) >= 3:
                            new_text = f"{parts[0]}Voor{parts[1]}Voor {client_name}{parts[2]}"
                            para.text = new_text
                
                # Pre-fill contact person name in "Naam" field
                if 'Naam' in text and contact_person and 'Opleidingsadviseur' in text:
                    # Add contact person after Naam label
                    if '\t' in text:
                        # Preserve tab structure
                        parts = text.split('Naam')
                        if len(parts) >= 2:
                            new_text = f"{parts[0]}Naam\t{contact_person}{parts[1]}"
                            para.text = new_text
                
                # Fill client company info in billing section
                if 'Jullie facturatiegegevens' in text:
                    # Next paragraphs should contain company address/VAT
                    # We'll handle this via placeholders in the field_map instead
                    pass
            
            # Search in tables for billing info
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            # Replace company address placeholder
                            if '{{client_company_address}}' in para.text and client_company_address:
                                self._replace_text_in_runs(para, '{{client_company_address}}', 
                                                          client_company_address)
                            # Replace VAT number placeholder
                            if '{{client_vat_number}}' in para.text and client_vat_number:
                                self._replace_text_in_runs(para, '{{client_vat_number}}', 
                                                          client_vat_number)
            
            print(f"[DOCX] Filled signature page with client: {client_name}, contact: {contact_person}")
            
        except Exception as e:
            print(f"Warning: Could not fill signature page: {str(e)}")
    
    def _replace_text_in_paragraph(self, paragraph, old_text: str, new_text: str):
        """
        Replace text in paragraph while preserving formatting
        """
        if old_text in paragraph.text:
            # Get full text
            full_text = paragraph.text
            # Replace
            new_full_text = full_text.replace(old_text, new_text)
            
            # Clear paragraph
            for run in paragraph.runs:
                run.text = ''
            
            # Add new text
            if paragraph.runs:
                paragraph.runs[0].text = new_full_text
            else:
                paragraph.add_run(new_full_text)
    
    def _create_new_document(
        self,
        offer_data: Dict[str, Any],
        crm_data: Dict[str, Any],
        training_data: Dict[str, Any],
        intake_data: Dict[str, Any]
    ) -> Document:
        """
        Create new document from scratch if template not available
        Uses Markdown formatting for content sections
        """
        doc = Document()
        
        # Add header
        self._add_header(doc, crm_data)
        
        # Add title
        title = doc.add_heading('OFFERTE', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add offer details
        doc.add_paragraph(f"Offertenummer: {self._generate_offer_number(crm_data, training_data)}")
        doc.add_paragraph(f"Datum: {self._format_dutch_date(crm_data.get('datum_offerte'))}")
        doc.add_paragraph(f"Klant: {crm_data.get('potentiele_klant', '')}")
        doc.add_paragraph(f"Contactpersoon: {crm_data.get('primaire_contact', '')}")
        doc.add_paragraph()
        
        # Add training title
        doc.add_heading(training_data.get('title', 'Opleiding'), level=2)
        doc.add_paragraph()
        
        # Add introduction with Markdown formatting
        doc.add_heading('Introductie', level=2)
        intro_text = offer_data.get('introduction', '')
        if intro_text:
            self._add_markdown_content(doc, intro_text)
        doc.add_paragraph()
        
        # Add objectives with Markdown formatting
        doc.add_heading('Opleidingsdoelstellingen', level=2)
        objectives_text = offer_data.get('training_objectives', '')
        if objectives_text:
            self._add_markdown_content(doc, objectives_text)
        doc.add_paragraph()
        
        # Add program overview with Markdown formatting
        doc.add_heading('Programmaoverzicht', level=2)
        program_text = offer_data.get('program_overview', '')
        if program_text:
            self._add_markdown_content(doc, program_text)
        doc.add_paragraph()
        
        # Add practical arrangements with Markdown formatting
        doc.add_heading('Praktische Regeling', level=2)
        practical_text = offer_data.get('practical_arrangements', '')
        if practical_text:
            self._add_markdown_content(doc, practical_text)
        doc.add_paragraph()
        
        # Add investment with Markdown formatting
        doc.add_heading('Investering', level=2)
        investment_text = offer_data.get('investment', '')
        if investment_text:
            self._add_markdown_content(doc, investment_text)
        doc.add_paragraph()
        
        # Add next steps with Markdown formatting
        doc.add_heading('Volgende Stappen', level=2)
        next_steps_text = offer_data.get('next_steps', '')
        if next_steps_text:
            self._add_markdown_content(doc, next_steps_text)
        doc.add_paragraph()
        
        # Add footer
        self._add_footer(doc)
        
        return doc
    
    def _add_header(self, doc: Document, crm_data: Dict[str, Any]):
        """
        Add header to document
        """
        section = doc.sections[0]
        header = section.header
        header_para = header.paragraphs[0]
        header_para.text = 'Syntra Bizz - In-company Opleidingen'
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    def _add_footer(self, doc: Document):
        """
        Add footer to document
        """
        section = doc.sections[0]
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.text = 'Syntra Antwerpen-Waasland vzw | www.syntra-ab.be'
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def _generate_offer_number(self, crm_data: Dict[str, Any], training_data: Dict[str, Any] = None) -> str:
        """
        Generate offer number in format: YYYY/NNN/INITIALS/ClientName TrainingType
        """
        if training_data is None:
            training_data = {}
        
        year = datetime.now().year
        client = crm_data.get('potentiele_klant', 'Client')
        training_title = training_data.get('title', '')
        
        # Clean names (keep spaces, remove special chars)
        client_clean = re.sub(r'[^\w\s-]', '', client).strip()
        training_clean = re.sub(r'[^\w\s-]', '', training_title).strip() if training_title else ''
        
        # Generate sequential number (3 digits, zero-padded)
        seq = str(uuid.uuid4().int)[:3].zfill(3)
        
        # Get initials from advisor or default to 'MP'
        advisor = crm_data.get('advisor_name', crm_data.get('type', ''))
        if advisor:
            # Extract initials from name
            initials = ''.join([word[0].upper() for word in advisor.split() if word])
        else:
            initials = 'MP'
        
        # Format: YYYY/NNN/INITIALS/ClientName TrainingType
        if training_clean:
            return f"{year}/{seq}/{initials}/{client_clean} {training_clean}"
        else:
            return f"{year}/{seq}/{initials}/{client_clean}"
    
    def _generate_filename(self, crm_data: Dict[str, Any], training_data: Dict[str, Any] = None) -> str:
        """
        Generate filename for the offer document
        """
        if training_data is None:
            training_data = {}
        
        offer_number = self._generate_offer_number(crm_data, training_data)
        # Replace slashes with underscores for filename
        filename_safe = offer_number.replace('/', '_')
        return f"{filename_safe}.docx"
