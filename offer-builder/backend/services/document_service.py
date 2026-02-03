import io
import logging
from typing import Optional
from fastapi import UploadFile
import PyPDF2
from docx import Document

logger = logging.getLogger(__name__)


class DocumentService:
    def __init__(self):
        pass
    
    def extract_text_from_pdf(self, file_content: bytes) -> str:
        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text_content = []
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    text = page.extract_text()
                    if text:
                        text_content.append(text)
                    logger.info(f"[DOCUMENT] Extracted text from PDF page {page_num + 1}")
                except Exception as e:
                    logger.warning(f"[DOCUMENT] Failed to extract page {page_num + 1}: {str(e)}")
                    continue
            
            full_text = '\n\n'.join(text_content)
            logger.info(f"[DOCUMENT] Total PDF text extracted: {len(full_text)} characters")
            return full_text
            
        except Exception as e:
            logger.error(f"[DOCUMENT] Error extracting PDF: {str(e)}")
            raise Exception(f"Failed to extract PDF content: {str(e)}")
    
    def extract_text_from_docx(self, file_content: bytes) -> str:
        try:
            docx_file = io.BytesIO(file_content)
            doc = Document(docx_file)
            
            text_content = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text_content.append(para.text)
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content.append(cell.text)
            
            full_text = '\n\n'.join(text_content)
            logger.info(f"[DOCUMENT] Total DOCX text extracted: {len(full_text)} characters")
            return full_text
            
        except Exception as e:
            logger.error(f"[DOCUMENT] Error extracting DOCX: {str(e)}")
            raise Exception(f"Failed to extract Word document content: {str(e)}")
    
    def extract_text_from_txt(self, file_content: bytes) -> str:
        try:
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            
            for encoding in encodings:
                try:
                    text = file_content.decode(encoding)
                    logger.info(f"[DOCUMENT] Text file decoded with {encoding}: {len(text)} characters")
                    return text
                except UnicodeDecodeError:
                    continue
            
            raise Exception("Unable to decode text file with common encodings")
            
        except Exception as e:
            logger.error(f"[DOCUMENT] Error extracting TXT: {str(e)}")
            raise Exception(f"Failed to extract text file content: {str(e)}")
    
    async def extract_document(self, file: UploadFile) -> str:
        try:
            file_content = await file.read()
            filename = file.filename.lower()
            
            logger.info(f"[DOCUMENT] Processing file: {file.filename} ({len(file_content)} bytes)")
            
            if filename.endswith('.pdf'):
                return self.extract_text_from_pdf(file_content)
            elif filename.endswith(('.docx', '.doc')):
                return self.extract_text_from_docx(file_content)
            elif filename.endswith('.txt'):
                return self.extract_text_from_txt(file_content)
            else:
                raise Exception(f"Unsupported file type: {filename}")
                
        except Exception as e:
            logger.error(f"[DOCUMENT] Error extracting document: {str(e)}")
            raise
